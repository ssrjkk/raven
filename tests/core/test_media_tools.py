from __future__ import annotations

import asyncio
import sys
import types
from pathlib import Path
from typing import Any

import httpx
import pytest
from PIL import Image

import raven.tools.media as media
from raven.core.task_engine.tool_registry import ToolRegistry


@pytest.fixture()
def ws(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setenv("RAVEN_WORKSPACE", str(tmp_path))
    return tmp_path


def _make_png(path: Path) -> None:
    img = Image.new("RGBA", (10, 10), (255, 0, 0, 128))
    img.save(str(path), format="PNG")


def _make_pdf(path: Path, text: str = "Hello media tools") -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_textbox(fitz.Rect(50, 50, 545, 792), text, fontsize=8)
    doc.save(str(path))
    doc.close()


class _FakeResp:
    def __init__(self, data: dict[str, Any], error: Exception | None = None):
        self._data = data
        self._error = error

    def raise_for_status(self) -> None:
        if self._error:
            raise self._error

    def json(self) -> dict[str, Any]:
        return self._data


class _FakeClient:
    def __init__(
        self,
        post_result: _FakeResp,
        get_results: list[_FakeResp] | None = None,
        post_error: Exception | None = None,
    ):
        self._post_result = post_result
        self._get_results = get_results or []
        self._post_error = post_error
        self.get_calls = 0
        self.last_post_kwargs: dict[str, Any] = {}

    async def __aenter__(self) -> _FakeClient:
        return self

    async def __aexit__(self, *_: object) -> bool:
        return False

    async def post(self, *args: object, **kwargs: object) -> _FakeResp:
        self.last_post_kwargs = dict(kwargs)
        if self._post_error:
            raise self._post_error
        return self._post_result

    async def get(self, *args: object, **kwargs: object) -> _FakeResp:
        if self._get_results:
            result = self._get_results[min(self.get_calls, len(self._get_results) - 1)]
        else:
            result = self._post_result
        self.get_calls += 1
        return result


def _patch_httpx(monkeypatch: pytest.MonkeyPatch, client: _FakeClient) -> None:
    monkeypatch.setattr(httpx, "AsyncClient", lambda **_: client)


async def _noop_sleep(*_: object, **__: object) -> None:
    return None


def test_image_generate_dispatch_dalle(monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_dalle(prompt: str, size: str, quality: str, n: int) -> str:
        return "dalle-ok"

    monkeypatch.setattr(media, "_image_generate_dalle", fake_dalle)
    assert asyncio.run(media.image_generate("x", model="dall-e-3")) == "dalle-ok"


def test_image_generate_dispatch_sd(monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_sd(prompt: str, size: str) -> str:
        return "sd-ok"

    monkeypatch.setattr(media, "_image_generate_sd", fake_sd)
    assert asyncio.run(media.image_generate("x", model="stable-diffusion")) == "sd-ok"
    assert asyncio.run(media.image_generate("x", model="sdxl")) == "sd-ok"


def test_image_generate_unknown_model() -> None:
    assert "Unknown model" in asyncio.run(media.image_generate("x", model="foo"))


def test_dalle_no_key(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "")
    assert "OPENAI_API_KEY" in asyncio.run(media._image_generate_dalle("x", "1024x1024", "standard", 1))


def test_dalle_success(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"data": [{"url": "https://x/img.png"}]})))
    out = asyncio.run(media._image_generate_dalle("a cat", "512x512", "hd", 1))
    assert "![Generated Image 1](https://x/img.png)" in out


def test_dalle_no_urls(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"data": [{"b64_json": "z"}]})))
    assert "No image URLs" in asyncio.run(media._image_generate_dalle("x", "1024x1024", "standard", 1))


def test_dalle_error(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    _patch_httpx(
        monkeypatch,
        _FakeClient(_FakeResp({"data": []}, error=httpx.ConnectError("boom"))),
    )
    assert "DALL-E generation failed" in asyncio.run(media._image_generate_dalle("x", "1024x1024", "standard", 1))


def test_sd_no_token(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "")
    assert "REPLICATE_API_TOKEN" in asyncio.run(media._image_generate_sd("x", ""))


def test_sd_success(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    post = _FakeResp({"urls": {"get": "https://x/poll"}})
    done = _FakeResp({"status": "succeeded", "output": ["https://x/out.png"]})
    client = _FakeClient(post, get_results=[done])
    _patch_httpx(monkeypatch, client)
    out = asyncio.run(media._image_generate_sd("x", "512x768"))
    assert "![Generated Image 1](https://x/out.png)" in out
    sent = client.last_post_kwargs["json"]["input"]
    assert sent["width"] == 512
    assert sent["height"] == 768


def test_sd_default_size(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    post = _FakeResp({"urls": {"get": "https://x/poll"}})
    done = _FakeResp({"status": "succeeded", "output": ["https://x/out.png"]})
    client = _FakeClient(post, get_results=[done])
    _patch_httpx(monkeypatch, client)
    asyncio.run(media._image_generate_sd("x", ""))
    sent = client.last_post_kwargs["json"]["input"]
    assert sent["width"] == 1024


def test_sd_bad_size_ignored(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    post = _FakeResp({"urls": {"get": "https://x/poll"}})
    done = _FakeResp({"status": "succeeded", "output": ["https://x/out.png"]})
    client = _FakeClient(post, get_results=[done])
    _patch_httpx(monkeypatch, client)
    asyncio.run(media._image_generate_sd("x", "bogus"))
    sent = client.last_post_kwargs["json"]["input"]
    assert sent["width"] == 1024


def test_sd_no_output(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    post = _FakeResp({"urls": {"get": "https://x/poll"}})
    done = _FakeResp({"status": "succeeded", "output": []})
    _patch_httpx(monkeypatch, _FakeClient(post, get_results=[done]))
    assert "No output" in asyncio.run(media._image_generate_sd("x", ""))


def test_sd_failed(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    post = _FakeResp({"urls": {"get": "https://x/poll"}})
    done = _FakeResp({"status": "failed", "error": "nope"})
    _patch_httpx(monkeypatch, _FakeClient(post, get_results=[done]))
    assert "prediction failed" in asyncio.run(media._image_generate_sd("x", ""))


def test_sd_timeout(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    post = _FakeResp({"urls": {"get": "https://x/poll"}})
    processing = _FakeResp({"status": "processing"})
    _patch_httpx(monkeypatch, _FakeClient(post, get_results=[processing]))
    assert "timed out" in asyncio.run(media._image_generate_sd("x", ""))


def test_sd_httpx_error(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_REPLICATE_API_TOKEN", "rpt")
    monkeypatch.setattr(asyncio, "sleep", _noop_sleep)
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"urls": {"get": "x"}}, error=httpx.ConnectError("boom"))))
    assert "Stable Diffusion generation failed" in asyncio.run(media._image_generate_sd("x", ""))


def test_image_edit_success(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    out = ws / "out.jpg"
    result = asyncio.run(media.image_edit(str(src), resize="5x5", crop="1,1,4,4", rotate=90, flip="horizontal", format="jpeg", output=str(out)))
    assert "Image edited" in result
    assert "crop(1,1,4,4)" in result
    assert "resize(5x5)" in result
    assert "flip(horizontal)" in result
    assert "data:image/jpeg;base64," in result
    assert out.exists()


def test_image_edit_no_output_path(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    result = asyncio.run(media.image_edit(str(src), flip="vertical"))
    assert "Image edited" in result
    assert (ws / "in_edited.png").exists()


def test_image_edit_jpeg_default_output(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    result = asyncio.run(media.image_edit(str(src), format="jpeg"))
    assert "Image edited" in result
    assert (ws / "in_edited.jpg").exists()


def test_image_edit_pillow_missing(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import builtins

    real_import = builtins.__import__

    def _no_pil(name: str, *args: Any, **kwargs: Any) -> Any:
        if name.startswith("PIL"):
            msg = "PIL disabled"
            raise ImportError(msg)
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", _no_pil)
    src = ws / "in.png"
    _make_png(src)
    assert "Pillow not available" in asyncio.run(media.image_edit(str(src)))


def test_image_edit_invalid_crop_count(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    assert "crop expects" in asyncio.run(media.image_edit(str(src), crop="1,2"))


def test_image_edit_invalid_crop_rect(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    assert "Invalid crop rectangle" in asyncio.run(media.image_edit(str(src), crop="4,1,1,4"))


def test_image_edit_invalid_resize(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    assert "resize expects" in asyncio.run(media.image_edit(str(src), resize="abc"))


def test_image_edit_invalid_flip(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    assert "flip expects" in asyncio.run(media.image_edit(str(src), flip="sideways"))


def test_image_edit_unsupported_format(ws: Path) -> None:
    src = ws / "in.png"
    _make_png(src)
    assert "Unsupported output format" in asyncio.run(media.image_edit(str(src), format="ICO"))


def test_image_edit_not_an_image(ws: Path) -> None:
    src = ws / "notes.txt"
    src.write_text("hello", encoding="utf-8")
    assert "Image processing failed" in asyncio.run(media.image_edit(str(src)))


def test_document_parse_missing_file(ws: Path) -> None:
    assert "File not found" in asyncio.run(media.document_parse(str(ws / "nope.pdf")))


def test_document_parse_pdf(ws: Path) -> None:
    pdf = ws / "doc.pdf"
    _make_pdf(pdf, "Hello media tools")
    out = asyncio.run(media.document_parse(str(pdf), "1"))
    assert "Hello media tools" in out
    assert "Extracted" in out


def test_document_parse_pdf_all_pages(ws: Path) -> None:
    pdf = ws / "doc.pdf"
    _make_pdf(pdf)
    out = asyncio.run(media.document_parse(str(pdf)))
    assert "Document: doc.pdf" in out


def test_document_parse_pdf_range(ws: Path) -> None:
    pdf = ws / "doc.pdf"
    _make_pdf(pdf)
    out = asyncio.run(media.document_parse(str(pdf), "1-1"))
    assert "Document: doc.pdf" in out


def test_document_parse_docx(ws: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello docx")
    doc.save(str(ws / "doc.docx"))
    out = asyncio.run(media.document_parse(str(ws / "doc.docx")))
    assert "Hello docx" in out


def _install_fake_pptx(monkeypatch: pytest.MonkeyPatch) -> None:
    mod: Any = types.ModuleType("pptx")

    class Para:
        def __init__(self, text: str) -> None:
            self.text = text

    class Frame:
        def __init__(self, texts: list[str]) -> None:
            self.paragraphs = [Para(t) for t in texts]

    class Shape:
        def __init__(self, texts: list[str]) -> None:
            self._frame = Frame(texts)

        @property
        def has_text_frame(self) -> bool:
            return True

        @property
        def text_frame(self) -> Frame:
            return self._frame

    class Slide:
        def __init__(self, texts: list[str]) -> None:
            self.shapes = [Shape(texts)]

    class Presentation:
        def __init__(self, _: str) -> None:
            self.slides = [Slide(["Hello slide"]), Slide([])]

    mod.Presentation = Presentation
    monkeypatch.setitem(sys.modules, "pptx", mod)


def _install_fake_openpyxl(monkeypatch: pytest.MonkeyPatch, long_cell: bool = False) -> None:
    mod: Any = types.ModuleType("openpyxl")

    class Cell:
        def __init__(self, value: object) -> None:
            self.value = value

    class Sheet:
        def __init__(self) -> None:
            val: object = "x" * 4000 if long_cell else ""
            self._rows = [[Cell("a"), Cell(val), Cell(3)]]

        def iter_row(self) -> Any:
            return iter(self._rows)

    class Wb:
        def __init__(self) -> None:
            self.sheetnames = ["Sheet1"]
            self._sheet = Sheet()

        def __getitem__(self, name: str) -> Sheet:
            return self._sheet

        def close(self) -> None:
            return

    mod.load_workbook = lambda *a, **k: Wb()
    monkeypatch.setitem(sys.modules, "openpyxl", mod)


def test_document_parse_pptx(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _install_fake_pptx(monkeypatch)
    pptx = ws / "slides.pptx"
    pptx.write_bytes(b"fake")
    out = asyncio.run(media.document_parse(str(pptx)))
    assert "Hello slide" in out
    assert "2 slides" in out


def test_document_parse_xlsx(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _install_fake_openpyxl(monkeypatch)
    xlsx = ws / "book.xlsx"
    xlsx.write_bytes(b"fake")
    out = asyncio.run(media.document_parse(str(xlsx)))
    assert "a |  | 3" in out


def test_document_parse_unsupported(ws: Path) -> None:
    txt = ws / "notes.txt"
    txt.write_text("hi", encoding="utf-8")
    assert "Unsupported format" in asyncio.run(media.document_parse(str(txt)))


def test_document_parse_import_error(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def _boom(path: Path, pages: str = "") -> str:
        msg = "no fitz"
        raise ImportError(msg)

    monkeypatch.setattr(media, "_parse_pdf", _boom)
    pdf = ws / "doc.pdf"
    _make_pdf(pdf)
    assert "Missing dependency" in asyncio.run(media.document_parse(str(pdf)))


def test_document_parse_general_error(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def _boom(path: Path, pages: str = "") -> str:
        msg = "boom"
        raise RuntimeError(msg)

    monkeypatch.setattr(media, "_parse_pdf", _boom)
    pdf = ws / "doc.pdf"
    _make_pdf(pdf)
    assert "Failed to parse" in asyncio.run(media.document_parse(str(pdf)))


def test_parse_pdf_pages_out_of_range(ws: Path) -> None:
    pdf = ws / "doc.pdf"
    _make_pdf(pdf)
    out = media._parse_pdf(pdf, "5,7")
    assert "Extracted 0 chars" in out


def test_parse_pdf_truncated(ws: Path) -> None:
    pdf = ws / "doc.pdf"
    _make_pdf(pdf, "x" * 2500)
    out = media._parse_pdf(pdf)
    assert "truncated" in out


def test_parse_pptx_branch(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _install_fake_pptx(monkeypatch)
    pptx = ws / "slides.pptx"
    pptx.write_bytes(b"fake")
    out = media._parse_pptx(pptx)
    assert "Extracted" in out


def test_parse_xlsx_branch(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _install_fake_openpyxl(monkeypatch)
    xlsx = ws / "book.xlsx"
    xlsx.write_bytes(b"fake")
    out = media._parse_xlsx(xlsx)
    assert "Workbook: book.xlsx" in out


def test_parse_xlsx_truncated(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _install_fake_openpyxl(monkeypatch, long_cell=True)
    xlsx = ws / "book.xlsx"
    xlsx.write_bytes(b"fake")
    out = media._parse_xlsx(xlsx)
    assert "truncated" in out


def test_video_info_access_denied(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("RAVEN_WORKSPACE", str(tmp_path / "ws"))
    outside = tmp_path / "outside.mp4"
    outside.write_bytes(b"x")
    assert "Access denied" in asyncio.run(media.video_info(str(outside)))


def test_video_info_prefix_collision(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    ws_dir = tmp_path / "ws"
    monkeypatch.setenv("RAVEN_WORKSPACE", str(ws_dir))
    sibling = Path(str(ws_dir) + "evil")
    sibling.mkdir(exist_ok=True)
    clip = sibling / "clip.mp4"
    clip.write_bytes(b"x")
    try:
        assert "Access denied" in asyncio.run(media.video_info(str(clip)))
    finally:
        clip.unlink(missing_ok=True)
        sibling.rmdir()


def test_video_info_missing_file(ws: Path) -> None:
    assert "File not found" in asyncio.run(media.video_info(str(ws / "nope.mp4")))


def test_video_info_no_ffmpeg(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setitem(sys.modules, "ffmpeg", None)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "ffmpeg-python not installed" in asyncio.run(media.video_info(str(vid)))


def test_video_info_probe_failure(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    mod: Any = types.ModuleType("ffmpeg")

    async def probe(*_: object, **__: object) -> dict[str, Any]:
        msg = "probe fail"
        raise RuntimeError(msg)

    mod.probe = probe
    monkeypatch.setitem(sys.modules, "ffmpeg", mod)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "ffprobe failed" in asyncio.run(media.video_info(str(vid)))


def test_video_info_success(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    mod: Any = types.ModuleType("ffmpeg")
    probe_result = {
        "streams": [
            {"codec_name": "h264", "codec_type": "video", "width": 1920, "height": 1080, "r_frame_rate": "30/1"},
            {"codec_name": "aac", "codec_type": "audio", "sample_rate": "48000", "channels": 2},
            {"codec_name": "none", "codec_type": "data"},
        ],
        "format": {"duration": "10.0", "size": "1000", "bit_rate": "800", "format_name": "mov,mp4"},
    }

    async def probe(*_: object, **__: object) -> dict[str, Any]:
        return probe_result

    mod.probe = probe
    monkeypatch.setitem(sys.modules, "ffmpeg", mod)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_info(str(vid)))
    assert "Video: clip.mp4" in out
    assert "h264" in out
    assert "aac" in out
    assert "Stream 2: data none" in out


def _install_fake_ffmpeg_chain(monkeypatch: pytest.MonkeyPatch, writer: Any = None, probe_result: Any = None) -> None:
    mod: Any = types.ModuleType("ffmpeg")

    class Node:
        def __init__(self) -> None:
            self._out: str | None = None

        def filter(self, *a: object, **k: object) -> Node:
            return self

        def output(self, *a: object, **k: object) -> Node:
            if a and isinstance(a[0], str):
                self._out = a[0]
            return self

        def overwrite_output(self) -> Node:
            return self

        def run(self, *a: object, **k: object) -> tuple[str, str]:
            if writer and self._out:
                writer(self._out)
            return ("", "")

    def _input(*a: object, **k: object) -> Node:
        return Node()

    mod.input = _input
    if probe_result is not None:
        mod.probe = probe_result
    monkeypatch.setitem(sys.modules, "ffmpeg", mod)


def test_video_thumbnail_missing_file(ws: Path) -> None:
    assert "File not found" in asyncio.run(media.video_thumbnail(str(ws / "nope.mp4")))


def test_video_thumbnail_no_ffmpeg(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setitem(sys.modules, "ffmpeg", None)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "ffmpeg-python not installed" in asyncio.run(media.video_thumbnail(str(vid)))


def test_video_thumbnail_success(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        img = Image.new("RGB", (4, 4), (0, 128, 0))
        img.save(out, format="JPEG")

    _install_fake_ffmpeg_chain(monkeypatch, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_thumbnail(str(vid), time_sec=1.5, size="320x240"))
    assert "Thumbnail saved" in out
    assert "![Thumbnail](data:image/jpeg;base64," in out


def test_video_thumbnail_custom_output(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        img = Image.new("RGB", (4, 4), (0, 128, 0))
        img.save(out, format="JPEG")

    _install_fake_ffmpeg_chain(monkeypatch, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out_path = ws / "shots" / "thumb.jpg"
    out = asyncio.run(media.video_thumbnail(str(vid), output=str(out_path)))
    assert "Thumbnail saved" in out
    assert out_path.exists()


def test_video_thumbnail_bad_output_image(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        Path(out).write_bytes(b"not an image")

    _install_fake_ffmpeg_chain(monkeypatch, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_thumbnail(str(vid)))
    assert "Thumbnail saved" in out
    assert "![Thumbnail]" not in out


def test_video_thumbnail_extract_failure(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        msg = "boom"
        raise RuntimeError(msg)

    _install_fake_ffmpeg_chain(monkeypatch, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "Thumbnail extraction failed" in asyncio.run(media.video_thumbnail(str(vid)))


def test_video_transcribe_access_denied(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("RAVEN_WORKSPACE", str(tmp_path / "ws"))
    outside = tmp_path / "outside.mp4"
    outside.write_bytes(b"x")
    assert "Access denied" in asyncio.run(media.video_transcribe(str(outside)))


def test_video_transcribe_prefix_collision(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    ws_dir = tmp_path / "ws"
    monkeypatch.setenv("RAVEN_WORKSPACE", str(ws_dir))
    sibling = Path(str(ws_dir) + "evil")
    sibling.mkdir(exist_ok=True)
    clip = sibling / "clip.mp4"
    clip.write_bytes(b"x")
    try:
        assert "Access denied" in asyncio.run(media.video_transcribe(str(clip)))
    finally:
        clip.unlink(missing_ok=True)
        sibling.rmdir()


def test_video_transcribe_no_httpx(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    monkeypatch.setitem(sys.modules, "httpx", None)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "httpx not available" in asyncio.run(media.video_transcribe(str(vid)))


def test_video_transcribe_missing_file(ws: Path) -> None:
    assert "File not found" in asyncio.run(media.video_transcribe(str(ws / "nope.mp4")))


def test_video_transcribe_no_key(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "")
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "OPENAI_API_KEY" in asyncio.run(media.video_transcribe(str(vid)))


def test_video_transcribe_ffmpeg_failure(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")

    class _FakeProc:
        async def communicate(self) -> tuple[bytes, bytes]:
            return (b"", b"")

    async def _fake_exec(*args: object, **kwargs: object) -> _FakeProc:
        return _FakeProc()

    monkeypatch.setattr(asyncio, "create_subprocess_exec", _fake_exec)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "ffmpeg audio extraction failed" in asyncio.run(media.video_transcribe(str(vid)))


def test_video_transcribe_success(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")

    class _FakeProc:
        def __init__(self) -> None:
            self._cmd: tuple[object, ...] = ()

        async def communicate(self) -> tuple[bytes, bytes]:
            assert self._cmd
            wav = Path(str(self._cmd[-1]))
            wav.write_bytes(b"RIFF-wav")
            return (b"", b"")

    async def _fake_exec(*args: object, **kwargs: object) -> _FakeProc:
        proc = _FakeProc()
        proc._cmd = args
        return proc

    monkeypatch.setattr(asyncio, "create_subprocess_exec", _fake_exec)
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"text": "hello transcription", "duration": 4.0})))
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_transcribe(str(vid), language="en"))
    assert "hello transcription" in out
    assert "Model: whisper-1" in out
    assert "Duration: 4.0s" in out
    leftovers = list(Path(__import__("tempfile").gettempdir()).glob("raven_audio_*.wav"))
    assert leftovers == []


def test_video_transcribe_http_error(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")

    class _FakeProc:
        def __init__(self) -> None:
            self._cmd: tuple[object, ...] = ()

        async def communicate(self) -> tuple[bytes, bytes]:
            assert self._cmd
            wav = Path(str(self._cmd[-1]))
            wav.write_bytes(b"RIFF-wav")
            return (b"", b"")

    async def _fake_exec(*args: object, **kwargs: object) -> _FakeProc:
        proc = _FakeProc()
        proc._cmd = args
        return proc

    monkeypatch.setattr(asyncio, "create_subprocess_exec", _fake_exec)
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"text": ""}, error=httpx.ConnectError("boom"))))
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "Transcription failed" in asyncio.run(media.video_transcribe(str(vid)))


def _install_fake_ffmpeg_probe(monkeypatch: pytest.MonkeyPatch, duration: float, writer: Any = None) -> None:
    mod: Any = types.ModuleType("ffmpeg")

    class Node:
        def __init__(self) -> None:
            self._out: str | None = None

        def filter(self, *a: object, **k: object) -> Node:
            return self

        def output(self, *a: object, **k: object) -> Node:
            if a and isinstance(a[0], str):
                self._out = a[0]
            return self

        def overwrite_output(self) -> Node:
            return self

        def run(self, *a: object, **k: object) -> tuple[str, str]:
            if writer and self._out:
                writer(self._out)
            return ("", "")

    def _input(*a: object, **k: object) -> Node:
        return Node()

    def _probe(*_: object, **__: object) -> dict[str, Any]:
        return {"format": {"duration": str(duration)}}

    mod.input = _input
    mod.probe = _probe
    monkeypatch.setitem(sys.modules, "ffmpeg", mod)


def test_extract_frames_access_denied(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("RAVEN_WORKSPACE", str(tmp_path / "ws"))
    outside = tmp_path / "outside.mp4"
    outside.write_bytes(b"x")
    assert "Access denied" in asyncio.run(media.video_extract_frames(str(outside)))


def test_extract_frames_prefix_collision(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    ws_dir = tmp_path / "ws"
    monkeypatch.setenv("RAVEN_WORKSPACE", str(ws_dir))
    sibling = Path(str(ws_dir) + "evil")
    sibling.mkdir(exist_ok=True)
    clip = sibling / "clip.mp4"
    clip.write_bytes(b"x")
    try:
        assert "Access denied" in asyncio.run(media.video_extract_frames(str(clip)))
    finally:
        clip.unlink(missing_ok=True)
        sibling.rmdir()


def test_extract_frames_missing_file(ws: Path) -> None:
    assert "File not found" in asyncio.run(media.video_extract_frames(str(ws / "nope.mp4")))


def test_extract_frames_no_ffmpeg(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setitem(sys.modules, "ffmpeg", None)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "ffmpeg-python not installed" in asyncio.run(media.video_extract_frames(str(vid)))


def test_extract_frames_no_duration(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _install_fake_ffmpeg_probe(monkeypatch, duration=0)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "Could not determine video duration" in asyncio.run(media.video_extract_frames(str(vid)))


def test_extract_frames_output_dir_denied(tmp_path: Path, monkeypatch: pytest.MonkeyPatch, ws: Path) -> None:
    _install_fake_ffmpeg_probe(monkeypatch, duration=10)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    outside = tmp_path / ".." / "frames_out"
    assert "Access denied" in asyncio.run(media.video_extract_frames(str(vid), output_dir=str(outside)))


def test_extract_frames_output_dir_prefix_denied(tmp_path: Path, monkeypatch: pytest.MonkeyPatch, ws: Path) -> None:
    _install_fake_ffmpeg_probe(monkeypatch, duration=10)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    sibling = Path(str(tmp_path) + "evil")
    sibling.mkdir(exist_ok=True)
    try:
        assert "Access denied" in asyncio.run(media.video_extract_frames(str(vid), output_dir=str(sibling)))
    finally:
        sibling.rmdir()


def test_extract_frames_relative_output_dir(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        img = Image.new("RGB", (8, 8), (0, 0, 255))
        img.save(out, format="JPEG")

    _install_fake_ffmpeg_probe(monkeypatch, duration=10, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_extract_frames(str(vid), output_dir="frames"))
    assert "Extracted 2 frames" in out
    assert (ws / "frames").exists()


def test_extract_frames_preview_failure(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        if out.endswith("0000_0s.jpg"):
            img = Image.new("RGB", (8, 8), (0, 0, 255))
            img.save(out, format="JPEG")

    _install_fake_ffmpeg_probe(monkeypatch, duration=10, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_extract_frames(str(vid)))
    assert "Extracted 2 frames" in out


def test_extract_frames_success(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        img = Image.new("RGB", (8, 8), (0, 0, 255))
        img.save(out, format="JPEG")

    _install_fake_ffmpeg_probe(monkeypatch, duration=10, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    out = asyncio.run(media.video_extract_frames(str(vid), interval_sec=5, max_frames=10, size="640x480"))
    assert "Extracted 2 frames from clip.mp4" in out
    assert "![frame](data:image/jpeg;base64," in out


def test_extract_frames_extraction_failure(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    def writer(out: str) -> None:
        msg = "boom"
        raise RuntimeError(msg)

    _install_fake_ffmpeg_probe(monkeypatch, duration=10, writer=writer)
    vid = ws / "clip.mp4"
    vid.write_bytes(b"x")
    assert "Frame extraction failed" in asyncio.run(media.video_extract_frames(str(vid)))


def test_image_analyze_missing_file(ws: Path) -> None:
    assert "File not found" in asyncio.run(media.image_analyze(str(ws / "nope.png")))


def test_image_analyze_no_key(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "")
    img = ws / "in.png"
    _make_png(img)
    assert "OPENAI_API_KEY" in asyncio.run(media.image_analyze(str(img)))


def test_image_analyze_unsupported_ext(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    txt = ws / "notes.txt"
    txt.write_text("hi", encoding="utf-8")
    assert "Unsupported image format" in asyncio.run(media.image_analyze(str(txt)))


def test_image_analyze_success(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"choices": [{"message": {"content": "a red circle"}}]})))
    img = ws / "in.png"
    _make_png(img)
    assert asyncio.run(media.image_analyze(str(img), "What is this?")) == "a red circle"


def test_image_analyze_mime_variants(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    client = _FakeClient(_FakeResp({"choices": [{"message": {"content": "ok"}}]}))
    _patch_httpx(monkeypatch, client)
    jpg = ws / "in.jpg"
    Image.new("RGB", (4, 4), (255, 255, 255)).save(jpg, format="JPEG")
    asyncio.run(media.image_analyze(str(jpg)))
    sent = client.last_post_kwargs["json"]["messages"][0]["content"][1]["image_url"]["url"]
    assert sent.startswith("data:image/jpeg;base64,")

    gif = ws / "in.gif"
    Image.new("P", (4, 4)).save(gif, format="GIF")
    asyncio.run(media.image_analyze(str(gif)))
    sent = client.last_post_kwargs["json"]["messages"][0]["content"][1]["image_url"]["url"]
    assert sent.startswith("data:image/gif;base64,")

    webp = ws / "in.webp"
    Image.new("RGB", (4, 4), (255, 255, 255)).save(webp, format="WEBP")
    asyncio.run(media.image_analyze(str(webp)))
    sent = client.last_post_kwargs["json"]["messages"][0]["content"][1]["image_url"]["url"]
    assert sent.startswith("data:image/webp;base64,")


def test_image_analyze_no_choices(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"error": "boom"})))
    img = ws / "in.png"
    _make_png(img)
    assert "API response" in asyncio.run(media.image_analyze(str(img)))


def test_image_analyze_http_error(ws: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(media, "_OPENAI_API_KEY", "sk-test")
    _patch_httpx(monkeypatch, _FakeClient(_FakeResp({"choices": []}), post_error=httpx.ConnectError("boom")))
    img = ws / "in.png"
    _make_png(img)
    assert "Image analysis failed" in asyncio.run(media.image_analyze(str(img)))


def test_register_media_tools() -> None:
    registry = ToolRegistry()
    media.register_media_tools(registry)
    names = [spec.name for spec in registry.list()]
    assert "image_generate" in names
    assert "image_edit" in names
    assert "image_analyze" in names
    assert "document_parse" in names
    assert "video_transcribe" in names
    assert "audio_transcribe" in names
    assert "video_extract_frames" in names
    assert "video_info" in names
    assert "video_thumbnail" in names
