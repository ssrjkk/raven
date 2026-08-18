from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

import httpx
import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from raven.core.media_api import create_media_router


@pytest.fixture()
def client(tmp_path: Path) -> TestClient:
    app = FastAPI()
    app.include_router(create_media_router(tmp_path))
    return TestClient(app)


def _make_png(path: Path) -> None:
    from PIL import Image

    img = Image.new("RGBA", (10, 10), (255, 0, 0, 128))
    img.save(str(path), format="PNG")


def _make_pdf(path: Path) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello pdf")
    doc.save(str(path))
    doc.close()


def _make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello docx")
    doc.save(str(path))


class _FakeResp:
    def __init__(self, data: dict[str, Any], error: Exception | None = None):
        self._data = data
        self._error = error

    def raise_for_status(self) -> None:
        if self._error:
            raise self._error

    def json(self) -> dict[str, Any]:
        return self._data


class _FakeClient:
    def __init__(self, resp: _FakeResp, post_error: Exception | None = None):
        self._resp = resp
        self._post_error = post_error

    async def __aenter__(self) -> _FakeClient:
        return self

    async def __aexit__(self, *_: object) -> bool:
        return False

    async def post(self, *args: object, **kwargs: object) -> _FakeResp:
        if self._post_error:
            raise self._post_error
        return self._resp


def _patch_client(monkeypatch: pytest.MonkeyPatch, resp: _FakeResp, post_error: Exception | None = None) -> None:
    monkeypatch.setattr(httpx, "AsyncClient", lambda **_: _FakeClient(resp, post_error))


def test_generate_missing_key(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    resp = client.post("/api/media/generate", params={"prompt": "a cat"})
    assert resp.status_code == 400
    assert "OPENAI_API_KEY" in resp.json()["detail"]


def test_generate_success(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "sk-test")
    _patch_client(monkeypatch, _FakeResp({"data": [{"url": "https://example.com/img.png"}]}))
    resp = client.post("/api/media/generate", params={"prompt": "a cat", "size": "512x512"})
    assert resp.status_code == 200
    body = resp.json()
    assert body["url"] == "https://example.com/img.png"
    assert body["prompt"] == "a cat"
    assert body["size"] == "512x512"


def test_generate_no_urls(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "sk-test")
    _patch_client(monkeypatch, _FakeResp({"data": [{"b64_json": "abc"}]}))
    resp = client.post("/api/media/generate", params={"prompt": "x"})
    assert resp.status_code == 500
    assert "No image URLs" in resp.json()["detail"]


def test_generate_http_error(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "sk-test")
    _patch_client(monkeypatch, _FakeResp({"data": []}, error=httpx.HTTPStatusError("400", request=httpx.Request("POST", "/"), response=httpx.Response(400))))
    resp = client.post("/api/media/generate", params={"prompt": "x"})
    assert resp.status_code == 500


def test_generate_connect_error(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("OPENAI_API_KEY", "sk-test")
    _patch_client(monkeypatch, _FakeResp({"data": []}), post_error=httpx.ConnectError("boom"))
    resp = client.post("/api/media/generate", params={"prompt": "x"})
    assert resp.status_code == 500


def test_process_resize_crop_rotate_flip(client: TestClient, tmp_path: Path) -> None:
    img = tmp_path / "in.png"
    _make_png(img)
    resp = client.post(
        "/api/media/process",
        params={"filepath": str(img), "resize": "5x5", "crop": "1,1,4,4", "rotate": 90, "flip": "horizontal", "output_format": "jpeg", "quality": 50},
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["format"] == "jpeg"
    assert body["width"] == 5
    assert body["height"] == 5
    assert body["data_url"].startswith("data:image/jpeg;base64,")


def test_process_default_format_vertical_flip(client: TestClient, tmp_path: Path) -> None:
    img = tmp_path / "in.png"
    _make_png(img)
    resp = client.post("/api/media/process", params={"filepath": str(img), "flip": "vertical"})
    assert resp.status_code == 200
    body = resp.json()
    assert body["format"] == "png"
    assert body["width"] == 10
    assert body["bytes"] > 0


def test_process_rejects_bad_ops(client: TestClient, tmp_path: Path) -> None:
    img = tmp_path / "in.png"
    _make_png(img)
    resp = client.post(
        "/api/media/process",
        params={"filepath": str(img), "crop": "a,b,c,d", "resize": "abc", "rotate": 0, "flip": "none"},
    )
    assert resp.status_code == 400
    resp = client.post("/api/media/process", params={"filepath": str(img), "resize": "0x100"})
    assert resp.status_code == 400
    resp = client.post("/api/media/process", params={"filepath": str(img), "crop": "1,2,1,2"})
    assert resp.status_code == 400


def test_process_bad_format_500(client: TestClient, tmp_path: Path) -> None:
    img = tmp_path / "in.png"
    _make_png(img)
    resp = client.post("/api/media/process", params={"filepath": str(img), "output_format": "bogus"})
    assert resp.status_code == 500


def test_process_traversal_403(client: TestClient, tmp_path: Path) -> None:
    outside = tmp_path / ".." / "outside.png"
    _make_png(outside)
    try:
        resp = client.post("/api/media/process", params={"filepath": str(outside)})
        assert resp.status_code == 403
    finally:
        outside.unlink(missing_ok=True)


def test_process_missing_file_404(client: TestClient, tmp_path: Path) -> None:
    resp = client.post("/api/media/process", params={"filepath": str(tmp_path / "nope.png")})
    assert resp.status_code == 404


def test_process_prefix_collision_403(client: TestClient, tmp_path: Path) -> None:
    sibling = Path(str(tmp_path) + "evil")
    sibling.mkdir(exist_ok=True)
    img = sibling / "in.png"
    _make_png(img)
    try:
        resp = client.post("/api/media/process", params={"filepath": str(img)})
        assert resp.status_code == 403
    finally:
        img.unlink(missing_ok=True)
        sibling.rmdir()


def test_parse_pdf(client: TestClient, tmp_path: Path) -> None:
    pdf = tmp_path / "doc.pdf"
    _make_pdf(pdf)
    resp = client.post("/api/media/parse", params={"filepath": str(pdf), "pages": "1"})
    assert resp.status_code == 200
    body = resp.json()
    assert "Hello pdf" in body["text"]
    assert body["metadata"]["pages"] == 1
    assert body["metadata"]["format"] == "pdf"


def test_parse_pdf_page_range(client: TestClient, tmp_path: Path) -> None:
    pdf = tmp_path / "doc.pdf"
    _make_pdf(pdf)
    resp = client.post("/api/media/parse", params={"filepath": str(pdf), "pages": "1-1"})
    assert resp.status_code == 200
    body = resp.json()
    assert "Hello pdf" in body["text"]
    assert body["metadata"]["extracted_pages"] == 1


def test_parse_docx(client: TestClient, tmp_path: Path) -> None:
    doc = tmp_path / "doc.docx"
    _make_docx(doc)
    resp = client.post("/api/media/parse", params={"filepath": str(doc)})
    assert resp.status_code == 200
    body = resp.json()
    assert "Hello docx" in body["text"]
    assert body["metadata"]["paragraphs"] == 1


def test_parse_pptx_missing_dep(client: TestClient, tmp_path: Path) -> None:
    pptx = tmp_path / "slides.pptx"
    pptx.write_bytes(b"placeholder")
    resp = client.post("/api/media/parse", params={"filepath": str(pptx)})
    assert resp.status_code == 500
    assert "Missing dependency" in resp.json()["detail"]


def test_parse_xlsx_missing_dep(client: TestClient, tmp_path: Path) -> None:
    xlsx = tmp_path / "book.xlsx"
    xlsx.write_bytes(b"placeholder")
    resp = client.post("/api/media/parse", params={"filepath": str(xlsx)})
    assert resp.status_code == 500
    assert "Missing dependency" in resp.json()["detail"]


def _install_fake_pptx(monkeypatch: pytest.MonkeyPatch) -> None:
    import types

    mod = types.ModuleType("pptx")

    class Para:
        def __init__(self, text: str) -> None:
            self.text = text

    class Frame:
        def __init__(self, texts: list[str]) -> None:
            self.paragraphs = [Para(t) for t in texts]

    class Shape:
        def __init__(self, texts: list[str]) -> None:
            self._frame = Frame(texts)

        @property
        def has_text_frame(self) -> bool:
            return True

        @property
        def text_frame(self) -> Frame:
            return self._frame

    class Slide:
        def __init__(self, texts: list[str]) -> None:
            self.shapes = [Shape(texts)]

    class Presentation:
        def __init__(self, _: str) -> None:
            self.slides = [Slide(["Hello slide"]), Slide([])]

    mod.Presentation = Presentation  # type: ignore[attr-defined]
    monkeypatch.setitem(sys.modules, "pptx", mod)


def _install_fake_openpyxl(monkeypatch: pytest.MonkeyPatch) -> None:
    import types

    mod = types.ModuleType("openpyxl")

    class Cell:
        def __init__(self, value: object) -> None:
            self.value = value

    class Sheet:
        def __init__(self) -> None:
            self._rows = [[Cell("a"), Cell(""), Cell(3)]]

        def iter_row(self):
            return iter(self._rows)

    class Wb:
        def __init__(self) -> None:
            self.sheetnames = ["Sheet1"]
            self._sheet = Sheet()

        def __getitem__(self, name: str) -> Sheet:
            return self._sheet

        def close(self) -> None:
            return

    mod.load_workbook = lambda *a, **k: Wb()  # type: ignore[attr-defined]
    monkeypatch.setitem(sys.modules, "openpyxl", mod)


def test_parse_pptx_branch(client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    _install_fake_pptx(monkeypatch)
    pptx = tmp_path / "slides.pptx"
    pptx.write_bytes(b"fake")
    resp = client.post("/api/media/parse", params={"filepath": str(pptx)})
    assert resp.status_code == 200
    body = resp.json()
    assert body["metadata"]["slides"] == 2
    assert "Hello slide" in body["text"]


def test_parse_xlsx_branch(client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    _install_fake_openpyxl(monkeypatch)
    xlsx = tmp_path / "book.xlsx"
    xlsx.write_bytes(b"fake")
    resp = client.post("/api/media/parse", params={"filepath": str(xlsx)})
    assert resp.status_code == 200
    body = resp.json()
    assert body["metadata"]["sheets"] == ["Sheet1"]
    assert "a |  | 3" in body["text"]


def test_parse_corrupt_pdf_500(client: TestClient, tmp_path: Path) -> None:
    bad = tmp_path / "broken.pdf"
    bad.write_bytes(b"not a real pdf")
    resp = client.post("/api/media/parse", params={"filepath": str(bad)})
    assert resp.status_code == 500


def test_parse_unsupported_format(client: TestClient, tmp_path: Path) -> None:
    txt = tmp_path / "notes.txt"
    txt.write_text("hello", encoding="utf-8")
    resp = client.post("/api/media/parse", params={"filepath": str(txt)})
    assert resp.status_code == 400
    assert "Unsupported format" in resp.json()["detail"]


def test_parse_missing_file_404(client: TestClient, tmp_path: Path) -> None:
    resp = client.post("/api/media/parse", params={"filepath": str(tmp_path / "nope.pdf")})
    assert resp.status_code == 404


def test_parse_traversal_403(client: TestClient, tmp_path: Path) -> None:
    outside = tmp_path / ".." / "outside.txt"
    outside.write_text("secret", encoding="utf-8")
    try:
        resp = client.post("/api/media/parse", params={"filepath": str(outside)})
        assert resp.status_code == 403
    finally:
        outside.unlink(missing_ok=True)


def test_parse_prefix_collision_403(client: TestClient, tmp_path: Path) -> None:
    sibling = Path(str(tmp_path) + "evil")
    sibling.mkdir(exist_ok=True)
    txt = sibling / "notes.txt"
    txt.write_text("secret", encoding="utf-8")
    try:
        resp = client.post("/api/media/parse", params={"filepath": str(txt)})
        assert resp.status_code == 403
    finally:
        txt.unlink(missing_ok=True)
        sibling.rmdir()


def test_upload_success(client: TestClient, tmp_path: Path) -> None:
    resp = client.post("/api/media/upload", files={"file": ("hello.png", b"\x89PNG-data", "image/png")})
    assert resp.status_code == 200
    body = resp.json()
    assert body["filename"] == "hello.png"
    assert body["size"] == 9
    assert (tmp_path / "hello.png").read_bytes() == b"\x89PNG-data"


def test_upload_no_filename(client: TestClient) -> None:
    body = (
        "--TESTBOUND\r\n"
        'Content-Disposition: form-data; name="file"; filename=""\r\n'
        "Content-Type: application/octet-stream\r\n"
        "\r\n"
        "data\r\n"
        "--TESTBOUND--\r\n"
    )
    resp = client.post(
        "/api/media/upload",
        content=body,
        headers={"Content-Type": "multipart/form-data; boundary=TESTBOUND"},
    )
    assert resp.status_code == 400


def test_upload_invalid_filename(client: TestClient) -> None:
    resp = client.post("/api/media/upload", files={"file": ("sub/..", b"data", "image/png")})
    assert resp.status_code == 400


def test_upload_write_error_500(client: TestClient, tmp_path: Path) -> None:
    dest_dir = tmp_path / "hello.png"
    dest_dir.mkdir()
    resp = client.post("/api/media/upload", files={"file": ("hello.png", b"data", "image/png")})
    assert resp.status_code == 500


def test_analyze(client: TestClient, monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    async def fake_analyze(filepath: str, prompt: str = "Describe this image in detail") -> str:
        assert prompt == "custom prompt"
        return "analyzed"

    monkeypatch.setattr("raven.tools.media.image_analyze", fake_analyze)
    img = tmp_path / "in.png"
    _make_png(img)
    resp = client.post("/api/media/analyze", params={"filepath": str(img), "prompt": "custom prompt"})
    assert resp.status_code == 200
    assert resp.json()["result"] == "analyzed"


def test_video_info(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_info(filepath: str) -> str:
        return "info!"

    monkeypatch.setattr("raven.tools.media.video_info", fake_info)
    resp = client.post("/api/media/video-info", params={"filepath": "clip.mp4"})
    assert resp.status_code == 200
    assert resp.json()["result"] == "info!"


def test_video_thumbnail(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_thumb(filepath: str, time_sec: float = 1.0, size: str = "320x240", output: str = "") -> str:
        assert time_sec == 2.5
        assert size == "640x480"
        return "thumb.png"

    monkeypatch.setattr("raven.tools.media.video_thumbnail", fake_thumb)
    resp = client.post("/api/media/video-thumbnail", params={"filepath": "clip.mp4", "time_sec": 2.5, "size": "640x480"})
    assert resp.status_code == 200
    assert resp.json()["result"] == "thumb.png"


def test_video_transcribe(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_transcribe(filepath: str, model: str = "whisper-1", language: str = "") -> str:
        assert model == "whisper-1"
        assert language == "en"
        return "transcript"

    monkeypatch.setattr("raven.tools.media.video_transcribe", fake_transcribe)
    resp = client.post("/api/media/video-transcribe", params={"filepath": "clip.mp4", "language": "en"})
    assert resp.status_code == 200
    assert resp.json()["result"] == "transcript"


def test_video_extract_frames(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    async def fake_frames(filepath: str, interval_sec: float = 5.0, max_frames: int = 10, size: str = "640x480") -> str:
        assert interval_sec == 1.0
        assert max_frames == 3
        assert size == "320x240"
        return "frames"

    monkeypatch.setattr("raven.tools.media.video_extract_frames", fake_frames)
    resp = client.post(
        "/api/media/video-extract-frames",
        params={"filepath": "clip.mp4", "interval_sec": 1.0, "max_frames": 3, "size": "320x240"},
    )
    assert resp.status_code == 200
    assert resp.json()["result"] == "frames"
